import os
import re
from typing import Dict, List, Tuple

SECTION_ALIASES = {
    "summary": ["summary", "objective", "profile"],
    "education": ["education", "academics", "academic background"],
    "experience": [
        "experience",
        "work experience",
        "professional experience",
        "employment",
    ],
    "projects": ["projects", "project experience"],
    "skills": ["skills", "technical skills", "technologies"],
    "certifications": ["certifications", "certificates"],
    "awards": ["awards", "honors", "achievements"],
}


def _sanitize_text(text: str) -> str:
    if not text:
        return ""
    normalized = text.replace("\u00a0", " ")
    normalized = normalized.replace("\u200b", "")
    normalized = normalized.replace("\u200c", "")
    normalized = normalized.replace("\u200d", "")
    normalized = normalized.replace("\u2018", "'").replace("\u2019", "'")
    normalized = normalized.replace("\u201c", '"').replace("\u201d", '"')
    normalized = normalized.replace("\u2022", "•")
    normalized = normalized.replace("\u2013", "–").replace("\u2014", "—")
    return normalized


def _strip_bullet(text: str) -> str:
    return re.sub(r"^[\s•\-\*]+", "", _sanitize_text(text)).strip()


def _detect_section(line: str) -> str | None:
    cleaned = re.sub(r"[:\-–—]+$", "", line.strip()).lower()
    if not cleaned:
        return None
    for section, aliases in SECTION_ALIASES.items():
        for alias in aliases:
            if cleaned == alias:
                return section
            if cleaned.startswith(alias) and len(cleaned) <= len(alias) + 5:
                return section
            if alias in cleaned and len(cleaned) <= 40:
                return section
    return None


def _extract_email(text: str) -> str:
    match = re.search(r"[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}", text, re.I)
    return match.group(0) if match else ""


def _extract_phone(text: str) -> str:
    match = re.search(
        r"(\+?\d[\d\s().-]{7,}\d)", text
    )
    return match.group(0) if match else ""


def _extract_urls(text: str) -> Dict[str, str]:
    urls = re.findall(r"https?://[^\s)]+", text)
    linkedin = next((u for u in urls if "linkedin.com" in u.lower()), "")
    github = next((u for u in urls if "github.com" in u.lower()), "")
    portfolio = ""
    for url in urls:
        lower = url.lower()
        if "linkedin.com" in lower or "github.com" in lower:
            continue
        portfolio = url
        break
    return {
        "linkedin_url": linkedin,
        "github_url": github,
        "portfolio_url": portfolio,
    }


def _extract_name(lines: List[str]) -> str:
    for line in lines[:5]:
        cleaned = line.strip()
        if cleaned and "@" not in cleaned and len(cleaned.split()) <= 6:
            return cleaned
    return ""


def _split_sections(lines: List[str]) -> Dict[str, List[str]]:
    section_positions: List[Tuple[str, int]] = []
    for idx, line in enumerate(lines):
        section = _detect_section(line)
        if section:
            section_positions.append((section, idx))

    section_positions.sort(key=lambda item: item[1])
    sections: Dict[str, List[str]] = {}
    for pos_idx, (section, start_idx) in enumerate(section_positions):
        end_idx = section_positions[pos_idx + 1][1] if pos_idx + 1 < len(section_positions) else len(lines)
        content = lines[start_idx + 1 : end_idx]
        sections[section] = content
    return sections


def _parse_entry_blocks(lines: List[str]) -> List[List[str]]:
    blocks: List[List[str]] = []
    current: List[str] = []
    for line in lines:
        if not line.strip():
            if current:
                blocks.append(current)
                current = []
            continue
        current.append(line.strip())
    if current:
        blocks.append(current)
    return blocks


def _parse_project_blocks(lines: List[str]) -> List[List[str]]:
    blocks: List[List[str]] = []
    current: List[str] = []
    for line in lines:
        stripped = line.strip()
        if not stripped:
            if current:
                blocks.append(current)
                current = []
            continue

        is_bullet = stripped.startswith(("•", "-", "*"))
        is_header = " | " in stripped and not is_bullet

        if is_header and current:
            blocks.append(current)
            current = [stripped]
            continue

        current.append(stripped)
    if current:
        blocks.append(current)
    return blocks


def _parse_entries(section_type: str, lines: List[str]) -> List[Dict[str, str]]:
    if not lines:
        return []

    if section_type == "skills":
        skills = " ".join(_strip_bullet(line) for line in lines).strip()
        if not skills:
            return []
        return [
            {
                "title": "",
                "organization": "",
                "location": "",
                "start_date": "",
                "end_date": "",
                "description": "",
                "technologies": skills,
                "order": 0,
            }
        ]

    blocks = _parse_project_blocks(lines) if section_type == "projects" else _parse_entry_blocks(lines)
    entries: List[Dict[str, str]] = []
    for order, block in enumerate(blocks):
        header = block[0]
        rest = block[1:]
        title = header
        organization = ""
        technologies = ""
        if " at " in header.lower():
            parts = re.split(r"\s+at\s+", header, flags=re.I)
            if len(parts) >= 2:
                title = parts[0].strip()
                organization = parts[1].strip()
        elif " - " in header:
            parts = header.split(" - ", 1)
            title = parts[0].strip()
            organization = parts[1].strip()
        elif " | " in header:
            parts = header.split(" | ", 1)
            title = parts[0].strip()
            if section_type == "projects":
                technologies = parts[1].strip()
            else:
                organization = parts[1].strip()

        description_lines = [_strip_bullet(line) for line in rest if line.strip()]
        description = ""
        if description_lines:
            description = "\n".join(f"• {line}" for line in description_lines if line)

        entries.append(
            {
                "title": title,
                "organization": organization,
                "location": "",
                "start_date": "",
                "end_date": "",
                "description": description,
                "technologies": technologies,
                "order": order,
            }
        )
    return entries


def parse_resume_text(text: str, filename: str = "") -> Dict[str, object]:
    cleaned_text = _sanitize_text(text)
    lines = [line.rstrip() for line in cleaned_text.splitlines()]
    non_empty_lines = [line.strip() for line in lines if line.strip()]

    email = _extract_email(text)
    phone = _extract_phone(text)
    urls = _extract_urls(text)
    full_name = _extract_name(non_empty_lines)

    sections = _split_sections(lines)
    summary_lines = sections.get("summary", [])
    summary = " ".join(_strip_bullet(line) for line in summary_lines).strip()

    parsed_sections = []
    for section_type in ["education", "experience", "projects", "skills"]:
        section_lines = sections.get(section_type, [])
        entries = _parse_entries(section_type, section_lines)
        parsed_sections.append(
            {
                "section_type": section_type,
                "section_title": section_type.replace("_", " ").title(),
                "entries": entries,
            }
        )

    resume_name = os.path.splitext(os.path.basename(filename or ""))[0].strip()
    if resume_name:
        resume_name = f"{resume_name} (Imported)"

    warnings = []
    if not full_name:
        warnings.append("Could not confidently detect a full name.")
    if not email:
        warnings.append("Could not find an email address.")

    return {
        "resume": {
            "name": resume_name,
            "full_name": full_name,
            "email": email,
            "phone": phone,
            "summary": summary,
            **urls,
        },
        "sections": parsed_sections,
        "warnings": warnings,
    }


def _extract_font_sizes_from_pdf(uploaded_file):
    """Extract font sizes from PDF to determine the dominant size."""
    try:
        import pdfplumber
        from collections import Counter
        
        uploaded_file.seek(0)  # Reset file pointer
        font_sizes = []
        
        with pdfplumber.open(uploaded_file) as pdf:
            for page in pdf.pages:
                chars = page.chars
                for char in chars:
                    if 'size' in char:
                        size = round(char['size'])
                        # Only consider reasonable body text sizes (8-14pt)
                        if 8 <= size <= 14:
                            font_sizes.append(size)
        
        if font_sizes:
            # Find most common font size
            size_counts = Counter(font_sizes)
            most_common = size_counts.most_common(1)[0][0]
            return most_common
        
        return 11  # Default fallback
    except Exception:
        return 11  # Default fallback


def parse_resume_file(uploaded_file) -> Dict[str, object]:
    filename = uploaded_file.name or ""
    ext = os.path.splitext(filename)[1].lower()
    
    base_font_size = 11  # Default

    if ext == ".pdf":
        import pdfplumber
        
        # Extract font size first
        base_font_size = _extract_font_sizes_from_pdf(uploaded_file)
        uploaded_file.seek(0)  # Reset for text extraction

        with pdfplumber.open(uploaded_file) as pdf:
            text = "\n".join(page.extract_text() or "" for page in pdf.pages)
    elif ext == ".docx":
        import docx

        doc = docx.Document(uploaded_file)
        
        # Try to extract font size from DOCX
        try:
            from collections import Counter
            font_sizes = []
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    if run.font.size:
                        # Font size in DOCX is in "twips" (1/20 of a point)
                        size_pt = run.font.size.pt
                        if 8 <= size_pt <= 14:
                            font_sizes.append(round(size_pt))
            if font_sizes:
                size_counts = Counter(font_sizes)
                base_font_size = size_counts.most_common(1)[0][0]
        except Exception:
            pass
        
        text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    elif ext == ".txt":
        raw = uploaded_file.read()
        try:
            text = raw.decode("utf-8")
        except UnicodeDecodeError:
            text = raw.decode("latin-1", errors="ignore")
    else:
        raise ValueError("Unsupported file type. Please upload a PDF, DOCX, or TXT file.")

    result = parse_resume_text(text, filename=filename)
    result['base_font_size'] = base_font_size
    return result
